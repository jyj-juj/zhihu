import base64
import hashlib
import hmac
import math
import os
import re
import sys
import threading
import time
import tkinter as tk
from http import cookiejar
from tkinter import filedialog
from urllib.parse import urlencode
import execjs
import requests
from PIL import Image
from docx import Document
from lxml import etree


class Sign():

    def __init__(self, username=None, password=None):
        self.username = username
        self.password = password
        self.e = {
            "client_id": "c3cef7c66a1843f8b3a9e6a1e3160e20",
            "grant_type": "password",
            "timestamp": "",
            "source": "com.zhihu.web",
            "signature": "",
            "username": self.username,
            "password": self.password,
            "captcha": "",
            "lang": "en",
            "utm_source": "",
            "ref_source": "other_https://www.zhihu.com/signin?next=%2F",
        }
        self.headers = {
            'accept-encoding': 'gzip, deflate, br',
            'host': 'www.zhihu.com',
            'referer': 'https://www.zhihu.com/',
            "user-agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/88.0.4324.104 Safari/537.36",
        }
        self.session = requests.session()
        self.session.headers = self.headers
        self.session.cookies = cookiejar.LWPCookieJar(filename='cookies.txt')

    # 获取验证码
    def get_capture(self):
        url = 'https://www.zhihu.com/api/v3/oauth/captcha?lang=en'
        b = self.session.get(url=url).json()
        if b["show_captcha"] == True:
            response = self.session.put(url=url)
            img_base64 = response.json()["img_base64"]
            img = base64.b64decode(img_base64)
            with open("img.jpg", "wb") as f:
                f.write(img)
            img = Image.open('img.jpg')
            img_thread = threading.Thread(target=img.show, daemon=True)
            img_thread.start()
            capt = input("请输入验证码:")
            self.e["captcha"] = capt
            self.session.post(url=url, data={'input_text': capt})

    # xsrf获取
    def get_xsrf(self):
        url = "https://www.zhihu.com/"
        self.session.get(url=url, allow_redirects=False)
        for c in self.session.cookies:
            if c.name == '_xsrf':
                _xsrf = c.value
                return _xsrf

    def timestamp(self):
        # 13位时间戳
        timestamp = int(time.time() * 1000)
        return timestamp

    def signature(self):
        self.e["timestamp"] = str(self.timestamp())
        ha = hmac.new(b'd1b964811afb40118a12068ff74a12f4', digestmod=hashlib.sha1)
        grant_type = self.e['grant_type']
        client_id = self.e['client_id']
        source = self.e['source']
        timestamp = self.e["timestamp"]
        ha.update(bytes((grant_type + client_id + source + timestamp), 'utf-8'))
        self.e["signature"] = ha.hexdigest()

    # formdata获取
    def formdata(self):
        self.signature()
        with open('formdata.js') as f:
            js = execjs.compile(f.read())
        e = urlencode(self.e)
        formdata = js.call('b', e)
        return formdata

    # 检查是否登录
    def check_sign(self):
        url = "https://www.zhihu.com/"
        response = self.session.get(url=url, allow_redirects=False)
        if response.status_code == 200:
            self.session.cookies.save()
            print("cookies保存成功！！！")
            return True
        else:
            return False

    # 检查cookies
    def check_cookies(self):
        try:
            # 加载cookies
            self.session.cookies.load(ignore_discard=True)
            #
            url = "https://www.zhihu.com/"
            response = self.session.get(url=url, allow_redirects=False)
            if response.status_code == 200:
                return True
            else:
                return False
        except FileNotFoundError:
            return "无cookie文件"

    def _input(self):
        self.username = input("请输入手机号:")
        self.password = input("请输入密码：")
        if "+86" not in self.username:
            self.username = "+86" + self.username

    # 知乎扫码登录
    def zhihu_png_sign(self):
        udid_url = 'https://www.zhihu.com/udid'
        qrcode_url = 'https://www.zhihu.com/api/v3/account/api/login/qrcode'
        qrcode_image_url = 'https://www.zhihu.com/api/v3/account/api/login/qrcode/{}/image'
        scaninfo_url = 'https://www.zhihu.com/api/v3/account/api/login/qrcode/{}/scan_info'
        response = self.session.post(url=udid_url)
        token_headers = {
            'Origin': 'https://www.zhihu.com',
            'Referer': 'https://www.zhihu.com/signup?next=%2F',
            'x-udid': response.content.decode('utf8')
        }
        token_headers.update(self.headers)
        response = self.session.post(url=qrcode_url, headers=token_headers)
        token = response.json()['token']
        # 下载保存二维码
        response = self.session.get(url=qrcode_image_url.format(token), headers=token_headers)
        img = response.content
        # 保存验证码
        with open("qrcode.jpg", "wb") as f:
            f.write(img)
        # 展示二维码
        img = Image.open('qrcode.jpg')
        img_thread = threading.Thread(target=img.show, daemon=True)
        img_thread.start()
        while True:
            response = self.session.get(url=scaninfo_url.format(token), headers=token_headers)
            text = response.text
            response_json = response.json()
            # --等待扫码/正在扫码
            if text == '{"status":0}':
                print('等待扫码')
                time.sleep(0.3)
            elif text == '{"status":1}':
                time.sleep(0.3)
                print('已扫码\n'
                      '请确认登录')
            elif "user_id" and "uid" in text:
                print("扫码成功！")
                break
            else:
                print("登入失败！")
                sys.exit(0)
        # hot_url = "https://www.zhihu.com/hot"
        # self.session.get(url = hot_url)
        # self.session.cookies.save()

    def account_password_login(self):
        self._input()
        self.e["username"] = self.username
        self.e["password"] = self.password
        self.get_capture()
        url = 'https://www.zhihu.com/api/v3/oauth/sign_in'
        formdata = self.formdata()
        headers = self.session.headers.copy()
        headers.update({
            'content-type': 'application/x-www-form-urlencoded',
            'x-zse-83': '3_2.0',
            'x-xsrftoken': self.get_xsrf()
        })
        self.session.post(url=url, data=formdata, headers=headers)

    def sign(self):
        b = self.check_cookies()
        if b == True:
            print("cookies有效！！！")
        else:
            if b == False:
                print("cookies失效！！！")
            else:
                print("没有cookie文件！！！")
            num = input("1代表账户密码登录\n"
                        "2代表扫码登录\n"
                        "请输入:")
            if num == "1":
                self.account_password_login()
            else:
                self.zhihu_png_sign()
            b = self.check_sign()
            if b == True:
                print("登录成功！！！")
                if num == "2":
                    if os.path.exists("qrcode.jpg"):
                        os.remove("qrcode.jpg")
                else:
                    if os.path.exists("img.jpg"):
                        os.remove("img.jpg")
            else:
                print("登入失败！！！")
        return self.session


'''打开选择文件夹对话框'''


def save_path():
    root = tk.Tk()
    root.withdraw()
    file_path = filedialog.askdirectory()  # 获得选择好的文件夹
    print(file_path)
    return file_path


class Spider():

    def __init__(self):
        self.sign = Sign()
        self.session = self.sign.sign()
        self.me_url_token = ""
        self.list_url_token = []
        self.article_info_dic = {}
        self.comment_list = []
        self.save_path = save_path()

    # 获取自己的url_token
    def get_me_url_token(self):
        url = "https://www.zhihu.com/api/v4/me"
        params = {
            "include": "allow_message,is_followed,is_following,is_org,is_blocking,employments,answer_count,follower_count,articles_count,gender,badge[?(type=best_answerer)].topics",
        }
        url = requests.get(url=url, params=params).url.replace("%28", "(").replace("%29", ")")
        response = self.session.get(url=url)
        dic = response.json()
        url_token = dic["url_token"]
        return url_token

    # 爬取
    def get_d_c0(self):
        cookies = requests.utils.dict_from_cookiejar(self.session.cookies)
        # print(cookies)
        # print(cookies["d_c0"])
        d_c0 = cookies["d_c0"]
        return d_c0

    # md5加密
    def md5(self, key):
        input_name = hashlib.md5()
        input_name.update(key.encode("utf-8"))
        return (input_name.hexdigest()).lower()

    # x_zse_86获取
    def x_zse_86(self, key):
        md5 = self.md5(key)
        with open('x-zse-86.js', 'r') as f:
            ctx1 = execjs.compile(f.read())
        encrypt_str = ctx1.call('b', md5)
        x_zse_86 = "2.0_" + encrypt_str
        return x_zse_86

    # 获取url_token
    def followees_url_token(self, page):
        base_url = "https://www.zhihu.com/api/v4/members/{}/followees".format(self.me_url_token)
        params = {
            "include": "data[*].answer_count,articles_count,gender,follower_count,is_followed,is_following,badge[?(type=best_answerer)].topics",
            "offset": page,
            "limit": "20",
        }
        response = requests.get(url=base_url, params=params)
        url = response.url
        url = url.replace("%2A", "*").replace("%28", "(").replace("%29", ")")
        key = '3_2.0+{}+{}'.format(url.split("https://www.zhihu.com")[1], self.get_d_c0())
        x_zse_86 = self.x_zse_86(key)
        headers = {
            "user-agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/88.0.4324.104 Safari/537.36",
            "x-zse-83": "3_2.0",
            "x-zse-86": x_zse_86,
        }
        self.session.headers = headers
        response = self.session.get(url=url, headers=headers)
        data = response.json()["data"]
        for i in data:
            url_token = i["url_token"]
            name = i["name"]
            # print(name,url_token)
            # (name,url_token)
            self.list_url_token.append((name, url_token))
        return data

    # 关注了，被关注数量以及文章数量
    def get_following_count(self, url_token):
        url = "https://www.zhihu.com/people/{}/following".format(url_token)
        response = self.session.get(url=url)
        content = response.content
        html = etree.HTML(content)
        following_list = html.xpath('//div[@class="NumberBoard-itemInner"]/strong/text()')
        articles_count = html.xpath('//*[@id="ProfileMain"]/div[1]/ul/li[5]/a/span/text()')[0]
        i, j = following_list
        # (关注了，关注者,文章数量)
        # 43,188这样数字的逗号要被去掉=====》》》j.replace(",","")
        following_count = (int(i), int(j.replace(",", "")), int(articles_count))
        return following_count

    def word(self):
        document = Document()
        document.add_heading(self.article_info_dic["title"], 0)  # 插入标题
        info = """
                名字：{}
                赞同数：{}
                评论数：{}
                文章链接：https://zhuanlan.zhihu.com/p/{}
        """.format(self.article_info_dic["author"], self.article_info_dic["comment_count"],
                   self.article_info_dic["voteup_count"], self.article_info_dic["id"])
        # print(info)
        document.add_paragraph(info)
        content = self.article_info_dic["content"].replace("<b>", "").replace("</b>", "")
        list = re.findall("(<p>.*?</p>)|(<figure.*?</figure>)|(<h2>.*?</h2>)", content, re.S)
        # print(list)
        for i in list:
            m, n, p = i
            q = m + n + p
            if "</p>" in q:
                q = q.replace("</p>", "").replace("<p>", "")
                if "</a>" in q:
                    href_name_list = re.findall('.*?href="(.*?)".*?>(.*?)</a>', q, re.S)
                    href, name = href_name_list[0]
                    q = "链接：{}，名字：{}".format(href, name)
                    document.add_paragraph(q)
                    # print(href, name)
                else:
                    document.add_paragraph(q)
                    # print(q)
            elif "</h2>" in q:
                q = q.replace("</h2>", "").replace("<h2>", "")
                document.add_heading(q)
                # print(q)
            elif "</figure>" in q:
                src_text_list = re.findall('.*?src="(.*?)".*?<figcaption>(.*?)</figcaption>.*?', q, re.S)
                if src_text_list != []:
                    src, text = src_text_list[0]
                    headers = {
                        "user-agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/88.0.4324.104 Safari/537.36",
                    }
                    response = requests.get(url=src, headers=headers)
                    content = response.content
                    with open("zhi_hu.png", "wb") as f:
                        f.write(content)
                    try:
                        document.add_picture('zhi_hu.png')
                    except:
                        pass
                    document.add_paragraph(text)
                    # print(text)
        document.add_heading("评论：")
        for comment in self.comment_list:
            document.add_paragraph(comment)
        print("{}/{}.docx".format(self.path, self.article_info_dic["title"]))
        document.save("{}/{}.docx".format(self.path, self.article_info_dic["title"]))

    def get_comment(self):
        self.comment_list = []
        base_url = "https://www.zhihu.com/api/v4/articles/{}/root_comments".format(self.article_info_dic["id"])
        print("评论一共{}条".format(self.article_info_dic["comment_count"]))
        page = 0
        while True:
            print("正在爬取评论{}页".format(page + 1))
            params = {
                "include": "data[*].article.column",
                "offset": str(page * 20),
                "limit": "20",
                "status": "open",
            }
            page = page + 1
            response = requests.get(url=base_url, params=params)
            url = response.url
            url = url.replace("%2A", "*").replace("%28", "(").replace("%29", ")")
            key = '3_2.0+{}+{}'.format(url.split("https://www.zhihu.com")[1], self.get_d_c0())
            x_zse_86 = self.x_zse_86(key)
            headers = {
                "referer": "https://www.zhihu.com/people/er-shi-si-86-26/following",
                "user-agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/88.0.4324.104 Safari/537.36",
                "x-zse-83": "3_2.0",
                "x-zse-86": x_zse_86,
            }
            self.session.headers = headers
            response = self.session.get(url=url, headers=headers)
            data = response.json()["data"]
            paging = response.json()["paging"]
            for i in data:
                label_list = ["</p>", "<p>", "<br>", "</a>"]
                content = i["content"]
                for label in label_list:
                    content = content.replace("{}".format(label), "")
                # print(content)
                name = i["author"]["member"]["name"]
                text = "{}:{}".format(name, content)
                self.comment_list.append(text)
                child_comments = i["child_comments"]
                if child_comments != []:
                    for j in child_comments:
                        name = j["author"]["member"]["name"]
                        content = j["content"]
                        # print(content)
                        text = "{}回复:{}".format(name, content)
                        self.comment_list.append(text)
            if paging["is_end"] == True:
                # print(self.comment_list)
                return ""

    def get_articles(self, url_token):
        print(url_token)
        base_url = "https://www.zhihu.com/api/v4/members/{}/articles".format(url_token)
        articles_num = self.get_following_count(url_token)[2]
        page = math.ceil(int(articles_num) / 20)
        print("一共{}页".format(page))
        for p in range(page):
            print("正在爬取第{}页...".format(p + 1))
            params = {
                "include": "data[*].comment_count,suggest_edit,is_normal,thumbnail_extra_info,thumbnail,can_comment,comment_permission,admin_closed_comment,content,voteup_count,created,updated,upvoted_followees,voting,review_info,is_labeled,label_info;data[*].author.badge[?(type=best_answerer)].topics",
                "offset": str(p * 20),
                "limit": "20",
                "sort_by": "voteups",
            }
            response = requests.get(url=base_url, params=params)
            url = response.url
            url = url.replace("%2A", "*").replace("%28", "(").replace("%29", ")")
            key = '3_2.0+{}+{}'.format(url.split("https://www.zhihu.com")[1], self.get_d_c0())
            x_zse_86 = self.x_zse_86(key)
            headers = {
                "user-agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/88.0.4324.104 Safari/537.36",
                "x-zse-83": "3_2.0",
                "x-zse-86": x_zse_86,
            }
            response = self.session.get(url=url, headers=headers)
            data = response.json()["data"]
            for i in data:
                voteup_count = i["voteup_count"]  # 赞同数量
                comment_count = i["comment_count"]  # 评论数量
                author = i["author"]["name"]  # 作者
                id = i["id"]  # 文章id
                title = i["title"]  # 文章标题
                content = i["content"]  # 文章内容

                article_info_dic = {}
                article_info_dic["voteup_count"] = voteup_count
                article_info_dic["comment_count"] = comment_count
                article_info_dic["author"] = author
                article_info_dic["id"] = id
                article_info_dic["comment_count"] = comment_count
                article_info_dic["title"] = title
                article_info_dic["content"] = content
                self.article_info_dic = article_info_dic

                if int(comment_count) != 0:
                    self.get_comment()
                self.word()

    def get_hot(self):
        time_path_title = str(int(time.time()))
        document = Document()
        document.add_heading('{} 热点'.format(time.strftime('%Y-%m-%d_%H:%M', time.localtime(time.time()))), 0)  # 插入标题
        base_url = "https://www.zhihu.com/api/v3/feed/topstory/hot-lists/total"
        params = {
            "limit": "50",
            "desktop": "true"
        }
        response = self.session.get(url=base_url, params=params)
        data = response.json()["data"]
        p = 1
        for i in data:
            print("正在爬取第{}条".format(p))
            p = p + 1
            title = i["target"]["title"]  # 标题
            metrics_area = i["detail_text"]  # 热度
            link = "https://www.zhihu.com/question/{}".format(i["target"]["id"])  # 文章链接
            img_url = i["children"][0]["thumbnail"].split("?source")[0]  # 文章封面
            document.add_heading(title, 0)  # 插入标题
            inf0 = """
            文章链接：{}
                    {}""".format(link, metrics_area)
            document.add_paragraph(inf0)
            if img_url != "":
                headers = {
                    "user-agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/88.0.4324.104 Safari/537.36",
                }
                response = requests.get(url=img_url, headers=headers)
                content = response.content
                with open("zhi_hu.png", "wb") as f:
                    f.write(content)
                try:
                    document.add_picture('zhi_hu.png')
                except:
                    pass

        document.save("{}/{}.docx".format(self.path, time_path_title))
        print("热点爬取成功")

    def articles_spider(self):
        # 获取自己的url_token
        self.me_url_token = self.get_me_url_token()
        # 获取自己关注了的数量
        following_count = self.get_following_count(self.me_url_token)
        page = math.ceil(following_count[0] / 20)
        # 把自己关注了的人的url_token，名字以tuple形式存入列表
        for p in range(page):
            self.followees_url_token(p * 20)
        # 把每个博主的文章进行爬取
        for i in self.list_url_token:
            name, url_token = i
            print("正在爬取{}".format(name))
            self.path = '{}/{}'.format(self.save_path, name)
            print(self.path)

            if os.path.exists(self.path) == False:
                os.mkdir(self.path)
            else:
                pass
            self.get_articles(url_token)

    def hot_spider(self):
        self.path = "{}/hot".format(self.save_path)
        if os.path.exists(self.path) == False:
            os.mkdir(self.path)
        else:
            pass
        self.get_hot()

    # 爬虫入口
    def spiders(self):
        print("1为文章爬取\n"
              "2为热点爬取")
        count = input("请输入数字：")
        # count = '2'
        if count == "1":
            self.articles_spider()
        else:
            self.hot_spider()


spider = Spider()
spider.spiders()
